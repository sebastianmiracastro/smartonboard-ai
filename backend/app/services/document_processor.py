import io
import re
from typing import List, Iterator
from pypdf import PdfReader
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

# Parámetros de fragmentación (contextos coherentes, no fragmentos arbitrarios).
# Fragmentos algo más grandes con buen solape: cada chunk conserva una idea completa
# y la recuperación (con vecinos) reconstruye el hilo sin cortar a la mitad.
DEFAULT_CHUNK_SIZE = 1000     # tamaño objetivo del contexto en caracteres
DEFAULT_OVERLAP = 180         # solape entre contextos (continuidad), recortado a palabra
MIN_CHUNK_LEN = 40            # descarta fragmentos demasiado cortos (encabezados, ruido)


# Numeración de página en sus formas típicas ("3", "Página 3", "3 / 10", "- 3 -").
_PAGE_NUM_RE = re.compile(r"^\s*(p[áa]g(?:ina)?\.?\s*)?\d+\s*(/\s*\d+)?\s*$", re.IGNORECASE)
_PAGE_DASH_RE = re.compile(r"^\s*[-–—]\s*\d+\s*[-–—]\s*$")


def _is_page_number(line: str) -> bool:
    return bool(_PAGE_NUM_RE.match(line) or _PAGE_DASH_RE.match(line))


def _strip_boilerplate(pages: List[str]) -> List[str]:
    """Quita encabezados/pies repetidos y la numeración de página de los PDFs.

    Los encabezados y pies (nombre de la empresa, título, 'Confidencial', 'Página N')
    se repiten en cada página y ensucian los fragmentos indexados. Se detectan las
    líneas CORTAS que aparecen en muchas páginas (cerca del inicio o del final) y se
    eliminan, junto con las líneas que son solo numeración."""
    from collections import Counter

    if len(pages) < 3:
        # Con pocas páginas no hay señal fiable de repetición; solo la numeración.
        return ["\n".join(l for l in p.splitlines() if not _is_page_number(l.strip()))
                for p in pages]

    per_page_lines = [[l.strip() for l in p.splitlines()] for p in pages]
    counter: Counter = Counter()
    for lines in per_page_lines:
        nonempty = [l for l in lines if l]
        # Candidatas a encabezado/pie: las 2 primeras y las 2 últimas líneas no vacías.
        # `set` para contar cada línea a lo sumo UNA vez por página (si no, una línea
        # de cuerpo en páginas cortas se contaría doble y se eliminaría por error).
        for l in set(nonempty[:2] + nonempty[-2:]):
            if 0 < len(l) <= 80:
                counter[l] += 1

    threshold = max(2, len(pages) // 2)
    boilerplate = {l for l, c in counter.items() if c >= threshold}

    out: List[str] = []
    for lines in per_page_lines:
        kept = [l for l in lines if l and l not in boilerplate and not _is_page_number(l)]
        out.append("\n".join(kept))
    return out


# Parámetros de OCR (para PDFs escaneados / con páginas en imagen).
OCR_MIN_CHARS = 15       # por debajo de esto, la página se considera sin capa de texto
OCR_DPI = 220            # resolución de render para OCR (nitidez vs. velocidad)
OCR_LANG = "spa+eng"     # idiomas de Tesseract (español + inglés)
# Extracción RICA: una página con imágenes/diagramas y POCO texto también se pasa por
# OCR para capturar el texto embebido en esas imágenes (que la capa de texto no ve).
IMAGE_OCR_MAX_TEXT = 900


def _merge_ocr(text_layer: str, ocr: str) -> str:
    """Añade al texto de la página las líneas de OCR que NO estén ya en la capa de
    texto (captura el texto dentro de imágenes/diagramas sin duplicar lo existente)."""
    base_low = (text_layer or "").lower()
    extra = [
        ln.strip() for ln in (ocr or "").splitlines()
        if len(ln.strip()) >= 4 and ln.strip().lower() not in base_low
    ]
    return (text_layer + "\n" + "\n".join(extra)).strip() if extra else text_layer


def _ocr_pixmap(pix) -> str:
    """OCR de una página renderizada con Tesseract. Degrada a '' si Tesseract o
    pytesseract no están instalados (no rompe la subida)."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        return pytesseract.image_to_string(img, lang=OCR_LANG)
    except Exception as e:
        print(f"OCR no disponible o falló: {e}")
        return ""


def _pdf_pages_pymupdf(file_bytes: bytes) -> List[str]:
    """Texto por página con PyMuPDF + OCR para una extracción RICA:

    - Páginas SIN capa de texto (escaneadas) → OCR completo de la página.
    - Páginas CON imágenes/diagramas y poco texto → OCR que se FUSIONA con la capa
      de texto para capturar lo que está embebido en las imágenes (sin duplicar).
    - Páginas de puro texto → se leen directo (rápido).
    """
    import fitz  # PyMuPDF
    pages: List[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            txt = page.get_text("text") or ""
            n = len(txt.strip())
            needs_ocr = n < OCR_MIN_CHARS
            # OCR-aumento: página con imágenes y texto escaso → seguramente hay texto
            # dentro de las imágenes que la capa de texto no capturó.
            augment = (not needs_ocr) and n < IMAGE_OCR_MAX_TEXT and bool(page.get_images())
            if needs_ocr or augment:
                try:
                    pix = page.get_pixmap(dpi=OCR_DPI, alpha=False)
                    ocr = _ocr_pixmap(pix)
                    if needs_ocr and len(ocr.strip()) > n:
                        txt = ocr
                    elif augment and ocr.strip():
                        txt = _merge_ocr(txt, ocr)
                except Exception as e:
                    print(f"No se pudo renderizar la página para OCR: {e}")
            pages.append(txt)
    return pages


def is_pdf_encrypted(file_bytes: bytes) -> bool:
    """True si el PDF está protegido con contraseña (no se puede abrir sin ella)."""
    try:
        import fitz
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            return bool(doc.needs_pass)
    except Exception:
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            return bool(getattr(reader, "is_encrypted", False))
        except Exception:
            return False


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracción SÓLIDA de PDF: PyMuPDF para la capa de texto (mejor que pypdf) y
    OCR (Tesseract) para páginas escaneadas o en imagen. Si PyMuPDF no está disponible
    por cualquier razón, cae a pypdf para no dejar de funcionar."""
    try:
        pages = _pdf_pages_pymupdf(file_bytes)
    except Exception as e:
        print(f"PyMuPDF no disponible, uso pypdf: {e}")
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [(page.extract_text() or "") for page in reader.pages]
    # Retira encabezados/pies y numeración repetidos antes de fragmentar.
    pages = _strip_boilerplate(pages)
    # Se separan las páginas con doble salto para conservar el límite de párrafo.
    return "\n\n".join(pages)


def _iter_block_items(parent) -> Iterator:
    """Itera párrafos Y tablas del cuerpo del documento en su ORDEN real.

    python-docx expone `doc.paragraphs` y `doc.tables` por separado y pierde el
    orden; recorriendo el XML del cuerpo se conserva la secuencia original."""
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _table_to_text(table: Table) -> str:
    """Convierte una tabla en líneas legibles 'celda | celda | …' (una por fila),
    para que el contenido de las tablas también se indexe y sea recuperable."""
    lines: List[str] = []
    for row in table.rows:
        # dict.fromkeys deduplica el texto repetido de celdas combinadas conservando el orden
        cells = list(dict.fromkeys(c.text.strip() for c in row.cells if c.text.strip()))
        if cells:
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extrae párrafos Y tablas del DOCX, en el orden en que aparecen.

    Antes solo se leían los párrafos, así que TODO el contenido en tablas (horarios,
    beneficios, escalas, contactos…) se perdía. Ahora también se extraen las tablas."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts: List[str] = []
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            if block.text.strip():
                parts.append(block.text)
        elif isinstance(block, Table):
            table_text = _table_to_text(block)
            if table_text:
                parts.append(table_text)
    return "\n".join(parts)


def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text(file_bytes: bytes, format: str) -> str:
    if format == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif format == "docx":
        return extract_text_from_docx(file_bytes)
    elif format == "txt":
        return extract_text_from_txt(file_bytes)
    return ""


# ─── LIMPIEZA Y FRAGMENTACIÓN ────────────────────────────────────────────────

def _clean(text: str) -> str:
    """Normaliza el texto preservando los límites de párrafo."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Une palabras cortadas por guion al final de línea: "infor-\nmación" -> "información"
    text = re.sub(r"-\n(\w)", r"\1", text)
    # Espacios/tabs múltiples -> uno solo (sin tocar los saltos de línea)
    text = re.sub(r"[ \t]+", " ", text)
    # Más de un salto -> separación de párrafo
    text = re.sub(r"\n{2,}", "\n\n", text)
    return text.strip()


def _split_sentences(text: str) -> List[str]:
    return [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]


def _trim_to_word(tail: str) -> str:
    """Recorta el solape para que empiece en una palabra completa."""
    tail = tail.lstrip()
    sp = tail.find(" ")
    return tail[sp + 1:] if sp != -1 else tail


def chunk_text(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_OVERLAP,
    min_len: int = MIN_CHUNK_LEN,
) -> List[str]:
    """Divide el texto en contextos coherentes.

    Estrategia: párrafos como unidad básica; los párrafos muy largos se parten por
    oraciones; las unidades se agrupan hasta `chunk_size` con un pequeño solape para
    no cortar ideas a la mitad. Así el agente recupera CONTEXTOS, no trozos sueltos.
    """
    text = _clean(text)
    if not text:
        return []

    # 1) Unidades: párrafos; los grandes se sub-dividen por oraciones
    units: List[str] = []
    for para in re.split(r"\n{2,}", text):
        para = re.sub(r"\s+", " ", para).strip()
        if not para:
            continue
        if len(para) <= chunk_size:
            units.append(para)
        else:
            acc = ""
            for sent in _split_sentences(para):
                if len(sent) > chunk_size:
                    # Oración monstruo: trocear por longitud como último recurso
                    if acc:
                        units.append(acc)
                        acc = ""
                    for i in range(0, len(sent), chunk_size):
                        units.append(sent[i:i + chunk_size])
                elif len(acc) + len(sent) + 1 <= chunk_size:
                    acc = (acc + " " + sent).strip()
                else:
                    if acc:
                        units.append(acc)
                    acc = sent
            if acc:
                units.append(acc)

    # 2) Agrupar unidades en contextos con solape por caracteres (recortado a palabra)
    chunks: List[str] = []
    current = ""
    for u in units:
        if not current:
            current = u
        elif len(current) + len(u) + 1 <= chunk_size:
            current = current + " " + u
        else:
            chunks.append(current.strip())
            tail = _trim_to_word(current[-overlap:]) if overlap > 0 else ""
            current = (tail + " " + u).strip() if tail else u
    if current.strip():
        chunks.append(current.strip())

    # 3) Descartar ruido (fragmentos demasiado cortos)
    return [c for c in chunks if len(c) >= min_len]
