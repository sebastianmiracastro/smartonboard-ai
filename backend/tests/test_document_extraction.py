"""Tests de extracción de texto de documentos (completitud del contenido)."""
import io

import pytest
from docx import Document as DocxDocument

from app.services.document_processor import (
    extract_text_from_docx, extract_text_from_pdf, chunk_text, _strip_boilerplate,
)


def _docx_bytes(build) -> bytes:
    d = DocxDocument()
    build(d)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_docx_extrae_parrafos_y_tablas_en_orden():
    def build(d):
        d.add_paragraph("Política de beneficios de la empresa")
        t = d.add_table(rows=3, cols=2)
        t.cell(0, 0).text = "Beneficio"
        t.cell(0, 1).text = "Detalle"
        t.cell(1, 0).text = "Vacaciones"
        t.cell(1, 1).text = "15 días hábiles al año"
        t.cell(2, 0).text = "Auxilio de transporte"
        t.cell(2, 1).text = "Según ley"
        d.add_paragraph("Para más información, contacta a RR.HH.")

    text = extract_text_from_docx(_docx_bytes(build))

    # Párrafos
    assert "Política de beneficios de la empresa" in text
    assert "contacta a RR.HH." in text
    # Contenido de la TABLA (antes se perdía por completo)
    assert "Vacaciones" in text and "15 días hábiles al año" in text
    assert "Auxilio de transporte" in text
    # Cada fila queda como una línea legible
    assert "Beneficio | Detalle" in text


def test_docx_tabla_se_indexa_en_chunks():
    def build(d):
        d.add_paragraph("Horarios de la compañía para el personal nuevo.")
        t = d.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "Horario de almuerzo del personal"
        t.cell(0, 1).text = "De 12:00 a 13:00, de lunes a viernes"

    text = extract_text_from_docx(_docx_bytes(build))
    chunks = chunk_text(text)
    joined = " ".join(chunks)
    assert "Horario de almuerzo del personal" in joined and "12:00 a 13:00" in joined


def _text_pdf_bytes(text: str) -> bytes:
    fitz = pytest.importorskip("fitz")
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text, fontsize=12)
    return doc.tobytes()


def _image_pdf_bytes(text: str) -> bytes:
    """PDF de una sola página que es una IMAGEN con texto (sin capa de texto)."""
    fitz = pytest.importorskip("fitz")
    from PIL import Image, ImageDraw, ImageFont
    img = Image.new("RGB", (1200, 300), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 110), text, fill="black", font=ImageFont.load_default(size=54))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    doc = fitz.open()
    page = doc.new_page(width=1200, height=300)
    page.insert_image(fitz.Rect(0, 0, 1200, 300), stream=buf.getvalue())
    return doc.tobytes()


def test_pdf_extrae_capa_de_texto():
    text = "Política de vacaciones: 15 días hábiles al año para todos los empleados."
    out = extract_text_from_pdf(_text_pdf_bytes(text))
    assert "vacaciones" in out.lower() and "15 días" in out


def test_pdf_escaneado_usa_ocr():
    pytesseract = pytest.importorskip("pytesseract")
    try:
        pytesseract.get_tesseract_version()
    except Exception:
        pytest.skip("Tesseract no está instalado en este entorno")
    out = extract_text_from_pdf(_image_pdf_bytes("VACACIONES 15 DIAS"))
    assert "VACACIONES" in out.upper()


def test_pdf_protegido_con_contrasena_se_detecta():
    fitz = pytest.importorskip("fitz")
    from app.services.document_processor import is_pdf_encrypted
    doc = fitz.open()
    doc.new_page()
    protegido = doc.tobytes(encryption=fitz.PDF_ENCRYPT_AES_256,
                            owner_pw="secreto", user_pw="secreto")
    assert is_pdf_encrypted(protegido) is True
    # Un PDF normal no está protegido.
    assert is_pdf_encrypted(_text_pdf_bytes("contenido con texto suficiente")) is False


def test_strip_boilerplate_quita_encabezados_y_paginacion():
    pages = [
        "ACME S.A. — Manual\nContenido único de la página uno.\nPágina 1",
        "ACME S.A. — Manual\nContenido distinto de la página dos.\nPágina 2",
        "ACME S.A. — Manual\nMás contenido en la página tres.\nPágina 3",
        "ACME S.A. — Manual\nCierre del documento en la página cuatro.\nPágina 4",
    ]
    joined = "\n".join(_strip_boilerplate(pages))
    # Encabezado repetido y numeración: fuera.
    assert "ACME S.A." not in joined
    assert "Página" not in joined
    # Cuerpo único de cada página: se conserva.
    assert "Contenido único de la página uno." in joined
    assert "Cierre del documento en la página cuatro." in joined
